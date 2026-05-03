"""Render a Markdown Data Science Memo into a formatted Word document.

Defaults are set to a safe, non-destructive regeneration target:
    - input:  archive/reports/DST_CRIME_ANALYSIS_MEMO_REFINED.md
    - output: reports/DST_CRIME_ANALYSIS_MEMO_REFINED_rebuilt.docx
    - figures directory root: figures/memo_refined

Override paths via CLI args to generate alternate memo versions.
"""

from __future__ import annotations

import re
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(".")
DEFAULT_MD_PATH = ROOT / "archive" / "reports" / "DST_CRIME_ANALYSIS_MEMO_REFINED.md"
DEFAULT_OUT_PATH = ROOT / "reports" / "DST_CRIME_ANALYSIS_MEMO_REFINED_rebuilt.docx"
DEFAULT_FIG_DIR = ROOT / "figures" / "memo_refined"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--md", type=Path, default=DEFAULT_MD_PATH, help="Input markdown file")
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT_PATH, help="Output .docx path")
    parser.add_argument(
        "--fig-dir",
        type=Path,
        default=DEFAULT_FIG_DIR,
        help="Base directory that image paths are resolved against",
    )
    return parser.parse_args()

args = parse_args()

MD_PATH = args.md
OUT_PATH = args.out
FIG_DIR = args.fig_dir

doc = Document()

# ── Page margins (1 inch all around) ─────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# ── Base font ─────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Heading styles ────────────────────────────────────────────────────────────
for lvl, size, bold in [(1, 16, True), (2, 13, True), (3, 11, True)]:
    h = doc.styles[f'Heading {lvl}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.bold = bold
    h.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# ── Helpers ───────────────────────────────────────────────────────────────────
def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


def apply_inline(run_parent, text):
    """Parse **bold**, *italic*, `code` inline markup and add runs."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        # plain text before match
        if m.start() > pos:
            run_parent.add_run(text[pos:m.start()])
        full = m.group(0)
        if full.startswith('**'):
            r = run_parent.add_run(m.group(2))
            r.bold = True
        elif full.startswith('*'):
            r = run_parent.add_run(m.group(3))
            r.italic = True
        else:
            r = run_parent.add_run(m.group(4))
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        run_parent.add_run(text[pos:])


def add_para(doc, text, style_name='Normal', align=None):
    p = doc.add_paragraph(style=style_name)
    if align:
        p.alignment = align
    apply_inline(p, text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_caption(doc, text):
    """Figure caption: italic, small, indented."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def add_image(doc, img_path, width=Inches(5.8)):
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run()
        run.add_picture(str(img_path), width=width)
    else:
        doc.add_paragraph(f'[Figure not found: {img_path.name}]')


def parse_table(lines):
    """Parse GFM table lines -> list of row lists."""
    rows = []
    for line in lines:
        if re.match(r'\s*\|[-| :]+\|\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows


def add_table_to_doc(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=n_cols)
    t.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = t.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            apply_inline(p, cell_text)
            p.paragraph_format.space_after = Pt(2)
            if i == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()  # spacing after table


# ── Parse & render ────────────────────────────────────────────────────────────
lines = MD_PATH.read_text(encoding='utf-8').splitlines()

i = 0
fig_bold_caption = None   # holds "Figure N — ..." bold line before <img>

while i < len(lines):
    line = lines[i]

    # Skip empty lines (but emit spacing via paragraph_format)
    if line.strip() == '':
        i += 1
        continue

    # Horizontal rule
    if re.match(r'^---+$', line.strip()):
        add_horizontal_rule(doc)
        i += 1
        continue

    # ATX headings
    m = re.match(r'^(#{1,3})\s+(.*)', line)
    if m:
        lvl = len(m.group(1))
        text = m.group(2).strip()
        doc.add_heading(text, level=lvl)
        i += 1
        continue

    # GFM table: collect all table lines
    if line.strip().startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        add_table_to_doc(doc, parse_table(table_lines))
        continue

    # HTML img tag — insert figure
    m = re.match(r'\s*<img\s+src="([^"]+)".*?alt="([^"]*)"', line)
    if m:
        src_raw = m.group(1)
        _alt = m.group(2)

        src_path = Path(src_raw)
        if src_path.is_absolute():
            img_path = src_path
        else:
            # Prefer resolving relative to the markdown file location.
            img_path = (MD_PATH.parent / src_path)
            if not img_path.exists():
                # Next, try interpreting the src as repo-root-relative.
                img_path = (ROOT / src_path)
            if not img_path.exists():
                # Fallbacks for custom fig-dir or simple filenames.
                img_path = (FIG_DIR / src_path)
            if not img_path.exists():
                img_path = (FIG_DIR / src_path.name)
        # Bold figure label was captured on previous non-empty pass
        if fig_bold_caption:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(fig_bold_caption)
            r.bold = True
            r.font.size = Pt(10)
            fig_bold_caption = None
        add_image(doc, img_path)
        i += 1
        continue

    # Bullet list item
    m = re.match(r'^[-*]\s+(.*)', line)
    if m:
        p = doc.add_paragraph(style='List Bullet')
        apply_inline(p, m.group(1))
        p.paragraph_format.space_after = Pt(3)
        i += 1
        continue

    # Numbered list item
    m = re.match(r'^\d+\.\s+(.*)', line)
    if m:
        p = doc.add_paragraph(style='List Number')
        apply_inline(p, m.group(1))
        p.paragraph_format.space_after = Pt(3)
        i += 1
        continue

    # Italic-only line (figure captions in *...*)
    if line.strip().startswith('*') and line.strip().endswith('*') and not line.strip().startswith('**'):
        caption_text = line.strip().strip('*')
        add_caption(doc, caption_text)
        i += 1
        continue

    # Bold figure label line: **Figure N — ...**
    m = re.match(r'^\*\*(.+?)\*\*\s*$', line.strip())
    if m:
        fig_bold_caption = m.group(1)
        i += 1
        continue

    # Regular paragraph
    add_para(doc, line.strip())
    i += 1

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
